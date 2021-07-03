"""
# -*- encoding: utf-8 -*-
# !/usr/bin/python3
@File:     generateSwagger.py
@Author:   shenfan
@Time:     2020/9/21 18:15
"""

from ruamel import yaml
import os
import json
from docx import Document
from docx.shared import Pt,RGBColor,Inches
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn


BASEDIR = os.path.join(os.path.dirname(os.path.dirname(__file__)),"file")
filetype = {
    "YAML":"Y",
    "WORD":"W"
}
class wordtemplate:
    def __init__(self):
        self.doc = Document()
        self.doc.styles['Normal'].font.name = u'微软雅黑'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        # 设置加粗
        self.doc.styles['Normal'].font.bold = False
        # 设置斜体
        self.doc.styles['Normal'].font.italic = False
        self.tablestyle = {
            "N":"Normal Table",
            "T":"Table Grid",
            "LS":"Light Shading",
            "LSA1":"Light Shading Accent 1",
            "LSA2":"Light Shading Accent 2",
            "LSA#3":"Light Shading Accent 3",
            "LSA4":"Light Shading Accent 4",
            "LSA5":"Light Shading Accent 5",
            "LSA6":"Light Shading Accent 6",
            "LL":"Light List",
            "LLA1":"Light List Accent 1",
            "LLA2":"Light List Accent 2",
            "LLA3":"Light List Accent 3",
            "LLA4":"Light List Accent 4",
            "LLA5": "Light List Accent 5",
            "LLA6": "Light List Accent 6",
            "LG":"Light Grid",
            "LGA1":"Light Grid Accent 1",
            "LGA2": "Light Grid Accent 2"
        }


    def paragraph(self,context):
        p = self.doc.add_paragraph()
        r = p.add_run(context)
        r.font.name = u"微软雅黑"
        r._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        r.font.size = Pt(10.5)

    def headline(self,context,level):
        r = self.doc.add_heading("", level).add_run(context)
        r.font.name = u"微软雅黑"
        r._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        if level == 2:
            r.font.size = Pt(22)
        if level == 3:
            r.font.size = Pt(16)
        if level == 4:
            r.font.size = Pt(14)



    def table(self,rows,cols,style,horizontalheader,*args):
        table=self.doc.add_table(rows,cols,style=style)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER | WD_TABLE_ALIGNMENT.LEFT | WD_TABLE_ALIGNMENT.RIGHT
        for row in range(rows):
            for col in range(cols):
                table.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        table.style.font.size = Pt(9.5)
        table.style.font.name = u"微软雅黑"
        # table.style.font.color.rgb = RGBColor(57, 89, 255)
        if horizontalheader:
            i = 0
            for arg in args:
                table.cell(0, i).text = arg
                i+=1
        else:
            i=0
            for arg in args:
                table.cell(i,0).text = arg
                i+=1
        return table

    def wordsave(self,file):
        self.doc.save(file)


class generateSwagger:
    def __init__(self,resoucefile,destinationfile,type):
        self.resoucefile = resoucefile
        self.destinationfile = destinationfile
        self.type = type

    def readfile(self):
        self.templist = []
        with open(self.resoucefile,"r",encoding="utf-8") as data:
            self.jsondata = json.load(data)
        for key0 in self.jsondata["paths"]:
            for key1 in self.jsondata["paths"][key0]:
                self.dic = {}
                self.dic.setdefault("path",key0)
                self.dic.setdefault("method",key1)
                self.dic.setdefault("tags",self.jsondata["paths"][key0][key1]["tags"][0])
                self.dic.setdefault("summary", self.jsondata["paths"][key0][key1]["summary"])
                self.dic.setdefault("operationId", self.jsondata["paths"][key0][key1]["operationId"])
                if "parameters" in self.jsondata["paths"][key0][key1]:
                    self.dic.setdefault("parameters", self.jsondata["paths"][key0][key1]["parameters"])
                self.templist.append(self.dic)

    def generateyaml(self):
        if not os.path.exists(self.destinationfile):
            open(self.destinationfile,"w").close()
        with open(self.destinationfile,"a+",encoding="utf-8") as data:
            yaml.dump(self.jsondata,data,Dumper=yaml.RoundTripDumper,allow_unicode=True)

    def generateword(self,topheader):
        self.wt = wordtemplate()
        self.wt.headline(topheader,level=2)
        for i in range(len(self.templist)):
            self.APIurl = self.templist[i].get("path")
            self.APIexample = "https://delivery.cbim.org.cn"+self.APIurl
            self.APImethod = self.templist[i].get("method")
            self.APIcomment = self.templist[i].get("summary")
            self.APIdescribetion = self.templist[i].get("summary") + self.templist[i].get("operationId")
            self.wt.headline(self.APIcomment,level=3)
            self.wt.paragraph("请手动添加接口描述")
            self.wt.headline("请求URL", level=4)
            tabletemplate_URL = self.wt.table(4,2,"Medium Grid 1 Accent 1",False,*("接口地址","接口示例","请求方式","接口备注"))
            self.info = [self.APIurl,self.APIexample,self.APImethod,self.APIdescribetion]
            for row in range(4):
                tabletemplate_URL.cell(row,1).text = self.info[row]
            self.wt.headline("请求头部参数",level=4)
            tabletemplate_Headerparams = self.wt.table(2,2,"Medium Grid 1 Accent 1",True,*("参数名称","参数说明"))
            tabletemplate_Headerparams.cell(1,0).text = "cookie"
            tabletemplate_Headerparams.cell(1,1).text = "SID=8EB02B52F4D9FCD12EAFDEAF3C811A68; tool.tk=0392447e0e9f898d98f20d650535d23a;"
            self.wt.headline("请求参数说明",level=4)
            paramsters = []
            if "parameters" in self.templist[i]:
                for j in range(len(self.templist[i]["parameters"])):
                    print("解析数据---->%s"%self.templist[i]["parameters"])
                    if 'schema' in self.templist[i]["parameters"][j] and "description" in self.templist[i]["parameters"][j]:
                        if "type" in self.templist[i]["parameters"][j]["schema"]:
                            paramster = [self.templist[i]["parameters"][j]["name"],self.templist[i]["parameters"][j]["schema"]["type"],self.templist[i]["parameters"][j]["description"],self.templist[i]["parameters"][j]["required"]]
                            paramsters.append(paramster)

                        else:
                            paramster = [self.templist[i]["parameters"][j]["name"],self.templist[i]["parameters"][j]["schema"],self.templist[i]["parameters"][j]["description"],self.templist[i]["parameters"][j]["required"]]
                            paramsters.append(paramster)
                            print("warning", self.templist[i]["path"],"-",self.templist[i]["method"],"*" * 10, "no type in schema and no type in parameters!!!")

                    elif 'schema' not in self.templist[i]["parameters"][j] and "description" in self.templist[i]["parameters"][j]:
                        paramster = [self.templist[i]["parameters"][j]["name"],self.templist[i]["parameters"][j]["type"],self.templist[i]["parameters"][j]["description"],self.templist[i]["parameters"][j]["required"]]
                        paramsters.append(paramster)
                        print("warning",self.templist[i]["path"],"-",self.templist[i]["method"], "*" * 10, "no schema in parameters!!!")

                    elif 'schema' in self.templist[i]["parameters"][j] and "description" not in self.templist[i]["parameters"][j]:
                        paramster = [self.templist[i]["parameters"][j]["name"],self.templist[i]["parameters"][j]["schema"]["type"],"-",self.templist[i]["parameters"][j]["required"]]
                        paramsters.append(paramster)
                        print("warning",self.templist[i]["path"],"-",self.templist[i]["method"], "*" * 10, "no description in parameters!!!")

                    else:
                        paramster = [self.templist[i]["parameters"][j]["name"],self.templist[i]["parameters"][j]["type"],"-",self.templist[i]["parameters"][j]["required"]]
                        paramsters.append(paramster)
                        print("warning", self.templist[i]["path"],"-",self.templist[i]["method"],"*" * 10, "no description in parameters!!!")
                print("解析数据--->%s"%paramsters)

            if "parameters" in self.templist[i]:
                num = len(self.templist[i]["parameters"])
                tabletemplate_Questparams = self.wt.table(num+1,4,"Medium Grid 1 Accent 1",True,*("参数名称","参数类型","参数说明","是否必须"))
                for row in range(num):
                    row += 1
                    for col in range(4):
                        print("解析数据---->%s"%paramsters[row-1][col])
                        tabletemplate_Questparams.cell(row,col).text = str(paramsters[row-1][col])
            else:
                pass
            self.wt.headline("请求参数示例",level=4)
            self.wt.paragraph("请手动添加参数示例")
            self.wt.headline("返回参数说明",level=4)
            tabletemplate_Responseparams = self.wt.table(4,2,"Medium Grid 1 Accent 1",True,*("返回参数名称","返回参数说明"))
            tabletemplate_Responseparams.cell(1,0).text = "code"
            tabletemplate_Responseparams.cell(1, 1).text = "接口返回状态码(0->成功;101->失败;201->没有权限)"
            tabletemplate_Responseparams.cell(2, 0).text = "message"
            tabletemplate_Responseparams.cell(2, 1).text = "接口返回状态信息(success：成功;error：失败;unauthorized：没有权限)"
            tabletemplate_Responseparams.cell(3, 0).text = "result"
            tabletemplate_Responseparams.cell(3, 1).text = "结果如下表所示"
            self.wt.headline("返回参数示例",level=4)
            self.wt.paragraph("手动添加接口示例")

        self.wt.wordsave(self.destinationfile)


if __name__ == "__main__":
    G = generateSwagger(os.path.join(BASEDIR,"api-docs-result.json"),os.path.join(BASEDIR,"doc.docx"),1)
    G.readfile()
    G.generateword("交付平台云端")
    # G.generateyaml()





